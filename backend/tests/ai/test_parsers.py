"""Tests for document parsers."""

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from enterprise_ai.ai.ingestion.parsers import DOCXParser, MarkdownParser, ParserFactory
from enterprise_ai.domain.value_objects.document import DocumentType


def test_markdown_parser(tmp_path: Path) -> None:
    file_path = tmp_path / "test.md"
    file_path.write_text("# Hello\n\nThis is a test document.", encoding="utf-8")
    parser = MarkdownParser()
    text = parser.parse(str(file_path))
    assert "Hello" in text
    assert "test document" in text


def test_docx_parser(tmp_path: Path) -> None:
    file_path = tmp_path / "test.docx"
    doc = DocxDocument()
    doc.add_paragraph("Enterprise policy document.")
    doc.add_paragraph("Second paragraph content.")
    doc.save(str(file_path))

    parser = DOCXParser()
    text = parser.parse(str(file_path))
    assert "Enterprise policy" in text
    assert "Second paragraph" in text


def test_parser_factory_from_extension() -> None:
    assert ParserFactory.get_parser(DocumentType.PDF) is not None
    assert ParserFactory.get_parser(DocumentType.MARKDOWN) is not None


def test_parse_markdown_from_bytes() -> None:
    content = b"# Title\n\nBody text here."
    text = ParserFactory.parse_from_bytes(DocumentType.MARKDOWN, content)
    assert "Title" in text
    assert "Body text" in text
