"""Document text extraction parsers."""

from abc import ABC, abstractmethod
from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from enterprise_ai.domain.exceptions import ValidationError
from enterprise_ai.domain.value_objects.document import DocumentType


class DocumentParser(ABC):
    """Extracts plain text from a document file."""

    @abstractmethod
    def parse(self, file_path: str) -> str:
        ...


class PDFParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        reader = PdfReader(file_path)
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(pages).strip()
        if not text:
            raise ValidationError("PDF contains no extractable text")
        return text


class DOCXParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        doc = DocxDocument(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paragraphs).strip()
        if not text:
            raise ValidationError("DOCX contains no extractable text")
        return text


class MarkdownParser(DocumentParser):
    def parse(self, file_path: str) -> str:
        text = Path(file_path).read_text(encoding="utf-8").strip()
        if not text:
            raise ValidationError("Markdown file is empty")
        return text


class ParserFactory:
    """Resolves the appropriate parser for a document type."""

    _parsers: dict[DocumentType, DocumentParser] = {
        DocumentType.PDF: PDFParser(),
        DocumentType.DOCX: DOCXParser(),
        DocumentType.MARKDOWN: MarkdownParser(),
    }

    @classmethod
    def get_parser(cls, document_type: DocumentType) -> DocumentParser:
        parser = cls._parsers.get(document_type)
        if parser is None:
            raise ValidationError(f"No parser for document type: {document_type}")
        return parser

    @classmethod
    def parse_from_bytes(cls, document_type: DocumentType, content: bytes) -> str:
        """Parse content from bytes (used in tests)."""
        if document_type == DocumentType.PDF:
            reader = PdfReader(BytesIO(content))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()
        if document_type == DocumentType.DOCX:
            doc = DocxDocument(BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
        return content.decode("utf-8").strip()
